from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import zipfile
import os

def unzip_file():
    cwd = os.getcwd()
    with zipfile.ZipFile(cwd+"/output/Compressed_images/compressor.zip", 'r') as zip_ref:
        zip_ref.extractall(cwd+"/output/Compressed_images")

def create_document_1(url :str):
    cwd = os.getcwd()
    document = Document()
    head = document.add_heading()
    head.add_run('WEBSITE ANALYSIS REPORT').bold = True
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.style = document.styles['Normal']
    font_head = head.style.font
    font_head.size = Pt(16)
    #font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
    head.paragraph_format.space_before = Pt(3)
    p1 = document.add_paragraph()
    p1.add_run("Website under analysis:- ").bold = True
    p1.add_run(url).italic = True
    p1.paragraph_format.space_after = Pt(60)


    p2 = document.add_paragraph()
    p2.add_run("Analyzer 1:- ").bold = True
    p2.add_run("GTmetrix").italic = True
    p2.add_run().add_break()
    p2.add_run("Results:  ").bold = True
    p2.add_run("GTmetrix  Grade  and  Web  Vitals").italic = True
    r1 = p2.add_run()
    r1.add_break()
    p2.paragraph_format.space_after = Pt(20)
    r1.add_picture(cwd+'/output/Gtmetrix/screenshot1.png', width=Inches(6.0))
    r1.add_break()
    p2.paragraph_format.space_after = Pt(20)
    r1.add_picture(cwd+'/output/Gtmetrix/screenshot2.png', width=Inches(6.0))
    p2.paragraph_format.space_after = Pt(60)

    p3 = document.add_paragraph()
    p3.add_run("Analyzer 2:- ").bold = True
    p3.add_run("Shopify  Speedbooster").italic = True
    p3.add_run().add_break()
    p3.add_run("Results:  ").bold = True
    p3.add_run("There  are  no  oversized  images  found").italic = True
    document.save(cwd+'/output/analysis_report_2.docx')


def create_document_2(url : str):
    cwd = os.getcwd()
    files = 0
    dir = cwd+"/output/oversized_images"
    for path in os.listdir(dir):
        if os.path.isfile(os.path.join(dir, path)):
            files += 1
    document = Document()
    head = document.add_heading()
    head.add_run('WEBSITE ANALYSIS REPORT').bold = True
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.style = document.styles['Normal']
    font_head = head.style.font
    font_head.size = Pt(16)
    head.paragraph_format.space_before = Pt(3)
    p1 = document.add_paragraph()
    p1.add_run("Website under analysis:- ").bold = True
    p1.add_run(url).italic = True
    p1.paragraph_format.space_after = Pt(60)


    p2 = document.add_paragraph()
    p2.add_run("Analyzer 1:- ").bold = True
    p2.add_run("GTmetrix").italic = True
    p2.add_run().add_break()
    p2.add_run("Results:  ").bold = True
    p2.add_run("GTmetrix  Grade  and  Web  Vitals").italic = True
    r1 = p2.add_run()
    r1.add_break()
    p2.paragraph_format.space_after = Pt(20)
    r1.add_picture(cwd+'/output/Gtmetrix/screenshot1.png', width=Inches(6.0))
    r1.add_break()
    p2.paragraph_format.space_after = Pt(20)
    r1.add_picture(cwd+'/output/Gtmetrix/screenshot2.png', width=Inches(6.0))
    p2.paragraph_format.space_after = Pt(60)


    p3 = document.add_paragraph()
    p3.add_run("Analyzer 2:- ").bold = True
    p3.add_run("Shopify  Speedbooster").italic = True
    p3.add_run().add_break()
    p3.add_run("Results:  ").bold = True
    p3.add_run("Oversized Images").italic = True
    r2 = p3.add_run()
    r2.add_break()
    p3.paragraph_format.space_after = Pt(20)
    for file in range(files):
        r2 = p3.add_run()
        r2.add_break()
        #p3.paragraph_format.space_after = Pt(20)
        r2.add_picture(cwd+'/output/oversized_images/image'+str(file+1)+'.jpeg', width=Inches(6.0))
        p3.paragraph_format.space_after = Pt(20)
    p3.paragraph_format.space_after = Pt(60)

    p5 = document.add_paragraph()
    p5.add_run("Image  Compressor:-  ").bold = True
    p5.add_run("https://compressor.io").italic = True
    p5.add_run().add_break()
    p5.add_run("Result:  ").bold = True
    p5.add_run("Total  bits  saved  after  Lossless  Compression  of  all  the  images").italic = True
    r4 = p5.add_run()
    r4.add_break()
    p5.paragraph_format.space_after = Pt(20)
    r4.add_picture(cwd+'/output/Compressed_Images/total_bits_saved.png', width=Inches(6.0))
    #r4.add_break()
    #p5.paragraph_format.space_after = Pt(20)
    #r4.add_picture(cwd+'/output/Compressed_Images/total_bits_saved.png', width=Inches(6.0))
    document.save(cwd+'/output/analysis_report_2.docx')
    print(files)

if __name__ == "__main__":
    url = "https://www.propero.in/"
    create_document_2(url)